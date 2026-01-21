from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from loguru import logger
from config import OUTPUT_DIR

class DOCXGenerator:
    def generate(self, contract_id: str, contract_type: str, content: str) -> str:
        try:
            doc = Document()
            title = doc.add_heading(f'{contract_type.replace("_", " ").title()} Contract', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for para in content.split('\n\n'):
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.style.font.size = Pt(11)
            
            file_path = OUTPUT_DIR / f"{contract_id}.docx"
            doc.save(str(file_path))
            logger.info(f"Generated DOCX: {file_path}")
            return str(file_path)
        except Exception as e:
            logger.error(f"Error generating DOCX: {e}")
            raise
    
    def get_file_path(self, contract_id: str) -> str:
        return str(OUTPUT_DIR / f"{contract_id}.docx")